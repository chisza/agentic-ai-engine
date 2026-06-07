"""Artifact tools for the summarizer agent.

Supports saving summaries as Markdown, PDF, or Word (.docx); and persisting
named summaries to the user-scoped artifact store so they are retrievable
across sessions.
"""

from __future__ import annotations

import io
import re

import structlog
from docx import Document
from fpdf import FPDF
from google.adk.tools.tool_context import ToolContext
from google.genai.types import Part

logger = structlog.get_logger(__name__)

_SUPPORTED_FORMATS = ("markdown", "pdf", "word")

_FORMAT_MIME: dict[str, tuple[str, str]] = {
    "markdown": ("text/markdown", ".md"),
    "pdf": ("application/pdf", ".pdf"),
    "word": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".docx",
    ),
}


def _to_markdown_bytes(text: str) -> bytes:
    return text.encode("utf-8")


_UNICODE_TO_ASCII = str.maketrans({
    "—": "--",   # em dash
    "–": "-",    # en dash
    "‘": "'",    # left single quote
    "’": "'",    # right single quote
    "“": '"',    # left double quote
    "”": '"',    # right double quote
    "…": "...",  # ellipsis
    "·": "-",    # middle dot
    "•": "-",    # bullet
})


def _to_pdf_bytes(text: str) -> bytes:
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_margins(15, 15, 15)
    pdf.set_font("Helvetica", size=11)
    # Strip markdown headers/bold/italic for clean PDF output
    clean = re.sub(r"#{1,6}\s+", "", text)
    clean = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", clean)
    clean = re.sub(r"_{1,2}(.+?)_{1,2}", r"\1", clean)
    # Replace Unicode typographic characters not supported by Helvetica (latin-1)
    clean = clean.translate(_UNICODE_TO_ASCII)
    clean = clean.encode("latin-1", errors="replace").decode("latin-1")
    pdf.multi_cell(0, 6, clean)
    return bytes(pdf.output())


def _to_docx_bytes(text: str) -> bytes:
    doc = Document()
    for line in text.split("\n"):
        stripped = line.strip()
        if stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped == "---":
            doc.add_paragraph("")
        else:
            doc.add_paragraph(stripped)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_artifact(text: str, fmt: str) -> tuple[Part, str, str]:
    """Return (Part, mime_type, extension) for the requested format."""
    mime, ext = _FORMAT_MIME[fmt]
    converters = {
        "markdown": _to_markdown_bytes,
        "pdf": _to_pdf_bytes,
        "word": _to_docx_bytes,
    }
    data = converters[fmt](text)
    part = Part(inline_data={"mime_type": mime, "data": data})
    return part, mime, ext


async def save_summary_as_artifact(
    summary: str,
    format: str,
    tool_context: ToolContext,
) -> dict:
    """Save a summary as a downloadable file in the session.

    Converts the provided summary text to the requested format and registers
    it as a session artifact so the UI can offer a download link.

    Args:
        summary: The clean summary text to save (not the full response including critic score).
        format: Output format — one of "markdown", "pdf", or "word".
    """
    fmt = format.lower().strip()
    if fmt not in _SUPPORTED_FORMATS:
        return {
            "status": "error",
            "message": f"Unsupported format '{format}'. Choose: {', '.join(_SUPPORTED_FORMATS)}.",
        }

    text = summary.strip()
    if not text:
        return {
            "status": "error",
            "message": "Summary text is empty.",
        }

    _, ext = _FORMAT_MIME[fmt]
    filename = f"summary{ext}"
    artifact, _, _ = _build_artifact(text, fmt)

    try:
        version = await tool_context.save_artifact(filename=filename, artifact=artifact)
        tool_context.state["last_saved_file"] = filename
        logger.info("Summary saved as session artifact", filename=filename, version=version)
        return {
            "status": "success",
            "message": f"Summary saved as '{filename}' (version {version}). It is now available for download.",
            "filename": filename,
        }
    except Exception as e:
        logger.error("Failed to save summary artifact", error=str(e))
        return {"status": "error", "message": str(e)}


async def save_summary_to_artifact_store(
    summary: str,
    name: str,
    format: str,
    tool_context: ToolContext,
) -> dict:
    """Persist a summary under a named key in the user artifact store.

    Unlike session artifacts, user-store artifacts survive across sessions and
    can be retrieved by name in any future conversation.

    Args:
        summary: The clean summary text to save (not the full response including critic score).
        name: A short identifier for this summary (e.g. "climate-report-2025").
              Must contain only letters, digits, hyphens, and underscores.
        format: Output format — one of "markdown", "pdf", or "word".
    """
    fmt = format.lower().strip()
    if fmt not in _SUPPORTED_FORMATS:
        return {
            "status": "error",
            "message": f"Unsupported format '{format}'. Choose: {', '.join(_SUPPORTED_FORMATS)}.",
        }

    if not re.match(r"^[\w\-]+$", name):
        return {
            "status": "error",
            "message": "Name may only contain letters, digits, hyphens, and underscores.",
        }

    text = summary.strip()
    if not text:
        return {
            "status": "error",
            "message": "Summary text is empty.",
        }

    _, ext = _FORMAT_MIME[fmt]
    # "user:" prefix tells GCS artifact service to store without a session_id —
    # making the artifact available in all future sessions for this user.
    filename = f"user:{name}{ext}"
    artifact, _, _ = _build_artifact(text, fmt)

    try:
        version = await tool_context.save_artifact(filename=filename, artifact=artifact)
        logger.info(
            "Summary saved to artifact store",
            store_name=name,
            filename=filename,
            version=version,
        )
        return {
            "status": "success",
            "message": (
                f"Summary '{name}' saved to your artifact store as '{filename}' "
                f"(version {version}). You can load it in any future session."
            ),
            "filename": filename,
            "store_name": name,
        }
    except Exception as e:
        logger.error("Failed to save to artifact store", error=str(e))
        return {"status": "error", "message": str(e)}


async def load_summary_from_artifact_store(
    name: str,
    tool_context: ToolContext,
) -> dict:
    """Load a previously stored summary from the user artifact store.

    Looks up an artifact saved via save_summary_to_artifact_store and returns
    its text content so you can present it to the user or continue working with it.

    Args:
        name: The identifier used when the summary was saved (e.g. "climate-report-2025").
              Optionally include the extension (e.g. "climate-report-2025.md").
    """
    # Try all extensions if name has none
    candidates: list[str] = []
    if "." in name and not name.startswith("user:"):
        candidates = [f"user:{name}"]
    else:
        base = name.lstrip("user:").split(".")[0]
        candidates = [f"user:{base}{ext}" for _, ext in _FORMAT_MIME.values()]

    for filename in candidates:
        try:
            part = await tool_context.load_artifact(filename=filename)
            if part is None:
                continue

            if part.inline_data and part.inline_data.data:
                raw = part.inline_data.data
                mime = part.inline_data.mime_type or ""
                if mime == "application/pdf" or filename.endswith(".pdf"):
                    return {
                        "status": "success",
                        "filename": filename,
                        "content": "[PDF artifact loaded — binary content, not displayable as text]",
                        "message": f"Artifact '{filename}' retrieved from store.",
                    }
                if (
                    mime.startswith("application/vnd.openxmlformats")
                    or filename.endswith(".docx")
                ):
                    return {
                        "status": "success",
                        "filename": filename,
                        "content": "[Word artifact loaded — binary content, not displayable as text]",
                        "message": f"Artifact '{filename}' retrieved from store.",
                    }
                text = raw.decode("utf-8", errors="replace")
                return {
                    "status": "success",
                    "filename": filename,
                    "content": text,
                    "message": f"Artifact '{filename}' retrieved from store.",
                }
            if part.text:
                return {
                    "status": "success",
                    "filename": filename,
                    "content": part.text,
                    "message": f"Artifact '{filename}' retrieved from store.",
                }
        except Exception as e:
            logger.warning("Error trying to load artifact", filename=filename, error=str(e))
            continue

    return {
        "status": "error",
        "message": (
            f"No artifact named '{name}' found in your store. "
            "Use list_artifact_store_contents to see what is available."
        ),
    }


async def list_artifact_store_contents(tool_context: ToolContext) -> dict:
    """List all summaries saved in the user artifact store.

    Returns the names and formats of all summaries previously persisted
    with save_summary_to_artifact_store, available in any session.
    """
    try:
        keys = await tool_context.list_artifacts()
        store_items = [k for k in keys if k.startswith("user:")]
        if not store_items:
            return {
                "status": "success",
                "items": [],
                "message": "Your artifact store is empty.",
            }
        return {
            "status": "success",
            "items": store_items,
            "message": f"Found {len(store_items)} item(s) in your artifact store.",
        }
    except Exception as e:
        logger.error("Failed to list artifact store", error=str(e))
        return {"status": "error", "message": str(e)}
